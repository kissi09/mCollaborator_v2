"""Builds the mCollaborator documentation DOCX from the Cyberteq template.

The template supplies the cover, the header/footer letterhead, the numbered
heading styles and the table style; everything after the table of contents is
replaced with the content in content_a.py / content_b.py.

    pip install python-docx
    python docgen.py
    .\\refresh-toc.ps1        # builds the table of contents through Word

The table of contents is a field. This script sets the document to refresh its
fields on open, so Word rebuilds it the first time anyone opens the file;
refresh-toc.ps1 does the same thing up front, which is what to run before the
document is sent anywhere or converted to PDF.
"""

import os
import re
import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

from content_a import PART_A
from content_b import PART_B

# docs/tools/docgen/docgen.py -> docs/
DOCS = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))

# The styling reference: a Cyberteq document whose cover, letterhead, heading
# numbering and table style this one inherits. Its own content is discarded.
TEMPLATE = os.path.join(DOCS, "mCollaborator Documentation.docx")
OUT = os.path.join(DOCS, "mCollaborator Extensive Documentation.docx")

BODY_FONT = "Wavehaus 42 Light"
BOLD_FONT = "Wavehaus 128 Bold"
CODE_FONT = "Consolas"
ORANGE = RGBColor(0xE9, 0x71, 0x32)
BULLET_NUMID = 44          # the template's own bullet list
LINE_SPACING = 360         # 1.5 lines, as every template paragraph uses


def set_line_spacing(p):
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), str(LINE_SPACING))
    spacing.set(qn("w:lineRule"), "auto")


def style_run(run, bold=False, code=False, size=12, color=None):
    run.font.name = CODE_FONT if code else (BOLD_FONT if bold else BODY_FONT)
    run.font.size = Pt(10 if code else size)
    if color is not None:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), run.font.name)


INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_rich_text(p, text, size=12):
    """Renders **bold** and `code` spans inside one paragraph."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            style_run(p.add_run(piece[2:-2]), bold=True, size=size)
        elif piece.startswith("`") and piece.endswith("`"):
            style_run(p.add_run(piece[1:-1]), code=True)
        else:
            style_run(p.add_run(piece), size=size)


class Builder:
    def __init__(self, doc, anchor):
        self.doc = doc
        self.anchor = anchor      # every new element is inserted after this one

    def _place(self, element):
        self.anchor.addnext(element)
        self.anchor = element

    def paragraph(self, style=None):
        p = self.doc.add_paragraph(style=style)
        self._place(p._p)
        set_line_spacing(p)
        return p

    def heading(self, text, level):
        p = self.paragraph(style="Heading %d" % level)
        for run in p.runs:
            run.text = ""
        r = p.add_run(text)
        r.font.size = Pt(14 if level == 2 else 12)
        rPr = r._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), BOLD_FONT)
        rPr.append(rFonts)
        return p

    def body(self, text):
        p = self.paragraph()
        add_rich_text(p, text)
        return p

    def bullet(self, text):
        p = self.paragraph(style="List Paragraph")
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), str(BULLET_NUMID))
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.insert(0, numPr)
        add_rich_text(p, text)
        return p

    def code(self, text):
        # Single-spaced and a size down: a diagram set 1.5-spaced at body size
        # wraps, and a wrapped diagram is not a diagram.
        for line in text.strip("\n").split("\n"):
            p = self.paragraph()
            pPr = p._p.get_or_add_pPr()
            spacing = pPr.find(qn("w:spacing"))
            spacing.set(qn("w:line"), "240")
            spacing.set(qn("w:after"), "0")
            run = p.add_run(line if line else " ")
            style_run(run, code=True)
            run.font.size = Pt(9)

    def table(self, headers, rows, widths=None):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.autofit = False
        tblPr = t._tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
        self._place(t._tbl)
        # Cells are single-spaced. The template's own tables are 1.5-spaced, but
        # they hold two wide columns; a five-column reference table set that way
        # turns every heading into a three-line tower.
        def compact(p):
            pPr = p._p.get_or_add_pPr()
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            spacing.set(qn("w:line"), "252")
            spacing.set(qn("w:lineRule"), "auto")
            spacing.set(qn("w:before"), "20")
            spacing.set(qn("w:after"), "20")

        for i, head in enumerate(headers):
            cell = t.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            compact(p)
            style_run(p.add_run(head), bold=True, color=ORANGE, size=11)
        for row in rows:
            cells = t.add_row().cells
            for i, value in enumerate(row):
                p = cells[i].paragraphs[0]
                compact(p)
                add_rich_text(p, str(value), size=10)
        # Column widths have to be set on every cell: Word reads the grid from
        # the cells, and a table left to autofit ignores the ones set on the
        # column object alone.
        if widths:
            grid = t._tbl.find(qn("w:tblGrid"))
            if grid is not None:
                for col, w in zip(grid.findall(qn("w:gridCol")), widths):
                    col.set(qn("w:w"), str(int(w * 1440)))
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Inches(w)
        # A spacer keeps the next block off the table's bottom rule.
        self.paragraph()
        return t

    def page_break(self):
        p = self.paragraph()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run = p.add_run()
        run._r.append(br)


def render(builder, blocks):
    for block in blocks:
        kind = block[0]
        value = block[1] if len(block) > 1 else None
        if kind == "h2":
            builder.heading(value, 2)
        elif kind == "h3":
            builder.heading(value, 3)
        elif kind == "p":
            builder.body(value)
        elif kind == "bullets":
            for item in value:
                builder.bullet(item)
        elif kind == "code":
            builder.code(value)
        elif kind == "table":
            headers, rows = value[0], value[1]
            widths = value[2] if len(value) > 2 else None
            builder.table(headers, rows, widths)
        elif kind == "pagebreak":
            builder.page_break()
        else:
            raise ValueError("unknown block: %r" % kind)


def main(argv=()):
    # Optional overrides, mostly so the tool can be run against a copy while the
    # real reference document is open in Word - Word holds an exclusive lock on
    # it, and python-docx reports that as "package not found".
    template = argv[0] if len(argv) > 0 else TEMPLATE
    out = argv[1] if len(argv) > 1 else OUT

    doc = docx.Document(template)
    body = doc.element.body
    children = list(body)

    # Keep the cover and the table of contents; drop the template's own content.
    toc_index = next(i for i, el in enumerate(children)
                     if el.tag == qn("w:sdt"))
    for el in children[toc_index + 1:]:
        if el.tag == qn("w:sectPr"):
            continue
        body.remove(el)

    builder = Builder(doc, children[toc_index])
    builder.page_break()
    render(builder, PART_A + PART_B)

    # Ask Word to rebuild the table of contents when the file is opened.
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)

    doc.save(out)
    print("wrote", out)


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
